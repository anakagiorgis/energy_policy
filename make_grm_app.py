import streamlit as st
from streamlit_chat import message as msg
import openai
import docx
import io

openai.api_key = "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"

st.title("Energy Policy Application")

st.write("***")

st.header("ChatGPT Turbo")

st.write("***")

if 'conversation_history' not in st.session_state:
	st.session_state.conversation_history = []

question = st.text_area("Ask a question regarding energy policy")
btn_q = st.button("Proceed")

if btn_q:
	st.session_state.conversation_history.append({"role": "user", "content": question})
	return_openai = openai.ChatCompletion.create(
		model = "gpt-3.5-turbo",
		messages = [{"role": "user", "content": question}],
		max_tokens = 3000,
		n=1
	)
	st.write(return_openai['choices'][0]['message']['content'])
	st.session_state.conversation_history.append(
		{"role": "assistant",
		"content": return_openai['choices'][0]['message']['content']})

if len(st.session_state.conversation_history) > 0:
	for i in range(len(st.session_state.conversation_history)):
		if (i%2) == 0:
			msg("Your Question: " + st.session_state.conversation_history[i]['content'], is_user=True)
		else:
			msg("AI Response: " + st.session_state.conversation_history[i]['content'])

if len(st.session_state.conversation_history) > 0:
	btn_s = st.button("Save Content")
	if btn_s:
		work = io.BytesIO()
		document = docx.Document()
		document.add_heading('Generated Content', level=1)
		for i in range(len(st.session_state.conversation_history)):
			if (i%2) == 0:
				document.add_heading('Question', level=2)
				document.add_paragraph(st.session_state.conversation_history[i]['content'])
			else:
				document.add_heading('Response', level=2)
				document.add_paragraph(st.session_state.conversation_history[i]['content'])
		document.save(work)
		st.download_button(label="Click Here to Save the Content",
				data=work,
				file_name="",
				mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

